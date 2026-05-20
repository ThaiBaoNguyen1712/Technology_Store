from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUTPUT_PATH = Path(
    r"D:\TaiLieuHocTap\MonHoc\Ky4\docx_render_review\Tech_Store_Docker_Setup_Guide.docx"
)
PROJECT_PATH = r"D:\Projects\Tech_Store"
RECOMMEND_PATH = r"D:\Projects\Tech_Store\services\recommend-api"


def set_cell_text(cell, text, bold=False, center=False):
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER if center else WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(text)
    run.bold = bold
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    run.font.size = Pt(11)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def shade_cell(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def add_heading(doc, text, level=1):
    p = doc.add_paragraph()
    p.style = f"Heading {level}"
    run = p.add_run(text)
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    if level == 1:
        run.font.size = Pt(15)
        run.font.color.rgb = RGBColor(31, 78, 121)
    elif level == 2:
        run.font.size = Pt(13)
    else:
        run.font.size = Pt(12)
    return p


def add_para(doc, text, bold_prefix=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.25
    if bold_prefix and text.startswith(bold_prefix):
        head, tail = bold_prefix, text[len(bold_prefix) :]
        run = p.add_run(head)
        run.bold = True
        run.font.name = "Times New Roman"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
        run.font.size = Pt(12)
        run = p.add_run(tail)
    else:
        run = p.add_run(text)
    for r in p.runs:
        r.font.name = "Times New Roman"
        r._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
        r.font.size = Pt(12)
    return p


def add_bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.line_spacing = 1.15
    run = p.add_run(text)
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    run.font.size = Pt(12)
    return p


def add_code(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.35)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    run.font.name = "Consolas"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Consolas")
    run.font.size = Pt(10.5)
    return p


def add_table(doc, headers, rows, widths=None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for i, header in enumerate(headers):
        set_cell_text(hdr[i], header, bold=True, center=True)
        shade_cell(hdr[i], "D9EAF7")
        if widths:
            hdr[i].width = widths[i]
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            set_cell_text(cells[i], value)
            if widths:
                cells[i].width = widths[i]
    doc.add_paragraph()
    return table


def configure_document(doc):
    section = doc.sections[0]
    section.page_width = Inches(8.27)
    section.page_height = Inches(11.69)
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.8)
    section.left_margin = Inches(0.9)
    section.right_margin = Inches(0.9)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Times New Roman"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    normal.font.size = Pt(12)

    for name in ["Title", "Heading 1", "Heading 2", "Heading 3"]:
        style = styles[name]
        style.font.name = "Times New Roman"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")


def build_doc():
    doc = Document()
    configure_document(doc)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("TÀI LIỆU TRIỂN KHAI VÀ CÀI ĐẶT HỆ THỐNG")
    run.bold = True
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    run.font.size = Pt(18)
    run.font.color.rgb = RGBColor(31, 78, 121)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run("Đồ án Tech_Store: ASP.NET Core MVC + Python Recommendation API")
    run.italic = True
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    run.font.size = Pt(13)

    add_para(doc, "Phiên bản tài liệu: 3.1", bold_prefix="Phiên bản tài liệu:")
    add_para(
        doc,
        "Mục đích: dùng làm tài liệu báo cáo đồ án và hướng dẫn cài đặt cho người đọc mới làm quen với dự án, kể cả khi trước đó chưa sử dụng Docker và chưa biết cách chạy ứng dụng web.",
        bold_prefix="Mục đích:",
    )

    add_heading(doc, "1. Phạm vi và mục tiêu tài liệu", 1)
    add_para(
        doc,
        "Tài liệu này trình bày lại từng bước theo cách dễ theo dõi, từ khâu chuẩn bị máy tính, cài phần mềm cần thiết, mở đúng thư mục dự án, chạy lệnh đầu tiên cho đến kiểm tra kết quả sau khi cài đặt."
    )
    add_bullet(doc, "Đối tượng sử dụng: sinh viên, giảng viên hướng dẫn, thành viên nhóm hoặc người mới tiếp nhận mã nguồn.")
    add_bullet(doc, "Giả định đầu vào: người đọc có thể chưa biết Docker là gì, chưa từng chạy ứng dụng ASP.NET Core hoặc Python API.")
    add_bullet(doc, "Mục tiêu triển khai: chỉ với một bộ source code, người đọc có thể làm theo tài liệu để chạy được cả website và recommendation service.")

    add_heading(doc, "2. Môi trường triển khai và đường dẫn dự án", 1)
    add_table(
        doc,
        ["Hạng mục", "Thông tin"],
        [
            ["Thư mục gốc dự án", PROJECT_PATH],
            ["Thư mục Recommendation API", RECOMMEND_PATH],
            ["File cấu hình Docker", rf"{PROJECT_PATH}\docker-compose.yml"],
            ["File biến môi trường", rf"{PROJECT_PATH}\.env"],
            ["Vị trí lưu tài liệu", str(OUTPUT_PATH)],
            ["Hệ điều hành kiểm thử", "Windows 11 64-bit"],
            ["Múi giờ máy kiểm thử", "Asia/Saigon"],
        ],
    )
    add_para(
        doc,
        "Nếu đặt dự án ở thư mục khác, người đọc chỉ cần nhớ một nguyên tắc đơn giản: mọi lệnh Docker trong tài liệu này phải được chạy tại đúng thư mục chứa file docker-compose.yml."
    )

    add_heading(doc, "3. Yêu cầu phần cứng", 1)
    add_table(
        doc,
        ["Thành phần", "Tối thiểu", "Khuyến nghị"],
        [
            ["CPU", "Core i3 / Ryzen 3, 2 nhân 4 luồng", "Core i5 / Ryzen 5 trở lên"],
            ["RAM", "8 GB", "16 GB trở lên"],
            ["Ổ đĩa trống", "15 GB", "25 GB trở lên"],
            ["Mạng", "Có Internet để tải image Docker", "Internet ổn định, không bị chặn registry"],
        ],
    )
    add_para(
        doc,
        "Các thông số khuyến nghị ở trên được đưa ra để việc build image và khởi động container diễn ra ổn định hơn. Nếu máy yếu hơn mức này thì hệ thống vẫn có thể chạy, nhưng thời gian khởi động có thể lâu và dễ phát sinh lỗi do thiếu bộ nhớ."
    )

    add_heading(doc, "4. Yêu cầu hệ điều hành và phần mềm", 1)
    add_table(
        doc,
        ["Hạng mục", "Yêu cầu bắt buộc / khuyến nghị"],
        [
            ["Hệ điều hành", "Windows 10 22H2 64-bit hoặc Windows 11 23H2/24H2 64-bit"],
            ["Docker Desktop", "Nên cài bản mới, hỗ trợ Docker Compose v2"],
            ["SQL Server", "SQL Server Express / Developer cài local trên máy host"],
            ["Công cụ SQL", "Khuyến nghị cài SSMS hoặc sqlcmd để kiểm tra DB"],
            [".NET SDK", "Khuyến nghị .NET 8 SDK nếu cần chạy dotnet ef ngoài Docker"],
            ["Quyền hệ thống", "Tài khoản Windows có quyền chạy Docker Desktop và kết nối SQL local"],
        ],
    )
    add_bullet(doc, "Không nên dùng Windows 32-bit hoặc các bản Windows quá cũ vì Docker Desktop và .NET 8 hoạt động không ổn định trên các môi trường này.")
    add_bullet(doc, "Nếu Docker Desktop báo thiếu điều kiện hệ thống, người dùng cần bật WSL2 hoặc các thành phần ảo hóa theo hướng dẫn của Docker.")

    add_heading(doc, "5. Cài đặt phần mềm nền tảng", 1)
    add_heading(doc, "5.1. Cài Docker Desktop", 2)
    add_bullet(doc, "Docker có thể hiểu đơn giản là công cụ giúp đóng gói sẵn môi trường chạy ứng dụng, nhờ đó người dùng không phải tự cài từng thành phần nhỏ bên trong project.")
    add_bullet(doc, "Tải Docker Desktop từ trang chính thức của Docker và cài đặt theo các bước mặc định.")
    add_bullet(doc, r"Vị trí cài mặc định thường là C:\Program Files\Docker\Docker.")
    add_bullet(doc, "Sau khi cài xong, cần mở Docker Desktop và chờ đến khi ứng dụng báo đã sẵn sàng. Nếu Docker chưa chạy, các lệnh phía sau sẽ không hoạt động.")
    add_code(doc, "docker --version")
    add_code(doc, "docker compose version")

    add_heading(doc, "5.2. Cài SQL Server local", 2)
    add_bullet(doc, "Có thể hiểu SQL Server là nơi lưu dữ liệu chính của hệ thống, ví dụ danh mục, sản phẩm, người dùng và các thông tin nghiệp vụ khác.")
    add_bullet(doc, "Có thể sử dụng SQL Server Developer hoặc Express tùy cấu hình máy.")
    add_bullet(doc, "Nên bật SQL Authentication và bảo đảm SQL Server có thể kết nối qua localhost:1433.")
    add_bullet(doc, "Nên cài thêm SSMS để việc mở database, restore dữ liệu và kiểm tra bảng dễ hơn đối với người mới.")
    add_code(doc, 'sqlcmd -S localhost -U sa -P "<mat_khau>" -C -Q "SELECT @@VERSION"')

    add_heading(doc, "5.3. Cài .NET 8 SDK nếu cần migrate thủ công", 2)
    add_bullet(doc, "Phần này chưa bắt buộc ngay từ đầu nếu người đọc chỉ muốn chạy thử hệ thống bằng Docker.")
    add_bullet(doc, "Chỉ cần cài .NET 8 SDK khi muốn dùng thêm lệnh dotnet ef để cập nhật database ngoài Docker.")
    add_code(doc, "dotnet --version")

    add_heading(doc, "6. Cấu trúc triển khai một bộ cài cho hai dịch vụ", 1)
    add_para(
        doc,
        "Trong cấu trúc hiện tại, mã nguồn Python Recommendation API đã được đưa vào chung repository Tech_Store tại thư mục services\\recommend-api. Vì vậy, người đọc không cần tải hai project rời nhau. Chỉ cần mở một project duy nhất rồi chạy docker compose là có thể khởi động các thành phần cần thiết."
    )
    add_bullet(doc, "Container web: cung cấp giao diện và logic ASP.NET Core MVC, chạy cổng 8080.")
    add_bullet(doc, "Container recommend-api: cung cấp API gợi ý sản phẩm FastAPI, chạy cổng 8000.")
    add_bullet(doc, "Container redis: lưu cache và dữ liệu trung gian phục vụ gợi ý.")
    add_bullet(doc, "SQL Server local: chạy ngoài Docker trên máy host, lưu database Electronics_Shop.")

    add_heading(doc, "7. Chuẩn bị file cấu hình .env", 1)
    add_para(doc, "File .env nằm tại thư mục gốc của dự án. Có thể hiểu đây là nơi khai báo các thông số cơ bản để container biết phải kết nối tới database nào và dùng tài khoản nào.")
    add_code(doc, "DB_HOST=host.docker.internal")
    add_code(doc, "DB_USER=sa")
    add_code(doc, "DB_PASSWORD=<mat_khau_sql>")
    add_code(doc, "APP_DB_NAME=Electronics_Shop")
    add_para(
        doc,
        "Trong cấu hình hiện tại, giá trị DB_HOST nên để là host.docker.internal. Tên này giúp container truy cập ngược lại SQL Server đang chạy trên máy Windows của người dùng."
    )

    add_heading(doc, "8. Chuẩn bị cơ sở dữ liệu và hướng dẫn migrate SQL", 1)
    add_heading(doc, "8.1. Phương án khuyến nghị cho đồ án hiện tại", 2)
    add_para(
        doc,
        "Với hiện trạng source code của đồ án, cách an toàn và dễ làm nhất là chuẩn bị sẵn database Electronics_Shop trên SQL Server local, hoặc restore một bản backup tương đương trước khi chạy Docker. Cách này thực tế hơn cho sinh viên mới tiếp cận dự án."
    )
    add_code(
        doc,
        'sqlcmd -S localhost -U sa -P "<mat_khau>" -C -Q "SET NOCOUNT ON; SELECT name FROM sys.databases WHERE name = \'Electronics_Shop\'"'
    )

    add_heading(doc, "8.2. Trường hợp tạo database mới", 2)
    add_code(
        doc,
        'sqlcmd -S localhost -U sa -P "<mat_khau>" -C -Q "CREATE DATABASE Electronics_Shop"'
    )
    add_para(
        doc,
        "Tuy nhiên, chỉ tạo database rỗng là chưa đủ để hệ thống chạy hoàn chỉnh. Trong phạm vi đồ án này, một số migration và dữ liệu nền chưa được tổ chức theo hướng khởi tạo hoàn toàn từ con số 0. Vì vậy, nếu người đọc là người mới, nên ưu tiên dùng database đã có sẵn thay vì cố dựng database trắng ngay từ đầu."
    )

    add_heading(doc, "8.3. Khi nào nên dùng lệnh migrate EF Core", 2)
    add_para(
        doc,
        "Lệnh migrate EF Core chỉ nên dùng khi nhóm đã có sẵn nền schema phù hợp hoặc đã restore trước một bản database mẫu. Trong báo cáo đồ án, nên trình bày trung thực rằng bước này cần có điều kiện dữ liệu nền phù hợp, thay vì khẳng định dự án luôn tạo được database hoàn chỉnh chỉ từ một lệnh migrate."
    )
    add_code(doc, "dotnet restore")
    add_code(
        doc,
        '$env:ConnectionStrings__DefaultConnection="Server=localhost,1433;Database=Electronics_Shop;User Id=sa;Password=<mat_khau>;TrustServerCertificate=True;Encrypt=False;"'
    )
    add_code(doc, "dotnet ef database update")

    add_heading(doc, "8.4. Kết luận phần SQL cho báo cáo", 2)
    add_para(
        doc,
        "Khi viết báo cáo đồ án, nên trình bày khiêm tốn và đúng với kết quả kiểm thử thực tế: Docker hiện được dùng để chạy web, recommend-api và Redis; còn SQL Server chạy local trên máy host với database Electronics_Shop đã tồn tại hoặc đã được restore trước."
    )

    add_heading(doc, "9. Các bước triển khai hệ thống bằng Docker", 1)
    add_bullet(doc, "Bước 1: mở thư mục dự án và xác định đúng thư mục gốc chứa file docker-compose.yml.")
    add_bullet(doc, rf"Bước 2: mở PowerShell tại thư mục {PROJECT_PATH}.")
    add_bullet(doc, "Bước 3: kiểm tra file .env đã điền đúng tài khoản SQL Server local.")
    add_bullet(doc, "Bước 4: bảo đảm Docker Desktop đang chạy và SQL Server local đã khởi động.")
    add_bullet(doc, "Bước 5: chạy lệnh bên dưới để Docker tự build và khởi động các container cần thiết.")
    add_code(doc, "docker compose up --build -d")
    add_bullet(doc, "Bước 6: sau khi lệnh chạy xong, kiểm tra trạng thái container.")
    add_code(doc, "docker compose ps")
    add_bullet(doc, "Bước 7: nếu cần xem bên trong hệ thống đang chạy ra sao, có thể mở log của từng service.")
    add_code(doc, "docker compose logs web -f")
    add_code(doc, "docker compose logs recommend-api -f")

    add_heading(doc, "10. Cách kiểm tra triển khai thành công", 1)
    add_para(doc, "Nếu người đọc chưa quen với cách chạy dự án, có thể kiểm tra theo các dấu hiệu đơn giản dưới đây. Khi các dấu hiệu này cùng xuất hiện, có thể hiểu là hệ thống đã chạy thành công ở mức cơ bản:")
    add_bullet(doc, "Lệnh docker compose ps hiển thị các container web, recommend-api và redis ở trạng thái Up hoặc healthy.")
    add_bullet(doc, "Mở trình duyệt và truy cập được website tại địa chỉ http://localhost:8080.")
    add_bullet(doc, "Truy cập được địa chỉ http://localhost:8000/health và nhận về mã phản hồi HTTP 200.")
    add_bullet(doc, "Log của web có dòng cho biết ứng dụng đang lắng nghe trên cổng 8080.")
    add_bullet(doc, "Log của recommend-api cho thấy tiến trình khởi động hoàn tất và đã nạp dữ liệu phục vụ gợi ý.")
    add_bullet(doc, "Khi mở website, các dữ liệu cơ bản như danh mục hoặc sản phẩm có thể hiển thị được thay vì trang trắng hoặc báo lỗi.")

    add_heading(doc, "11. Cách kiểm tra nhanh bằng lệnh", 1)
    add_code(doc, "Invoke-WebRequest http://localhost:8080 | Select-Object StatusCode")
    add_code(doc, "Invoke-WebRequest http://localhost:8000/health | Select-Object StatusCode")
    add_code(doc, "docker compose ps")

    add_heading(doc, "12. Các lệnh vận hành thường dùng", 1)
    add_code(doc, "docker compose up -d")
    add_code(doc, "docker compose down")
    add_code(doc, "docker compose restart web")
    add_code(doc, "docker compose restart recommend-api")
    add_code(doc, "docker compose logs redis -f")

    add_heading(doc, "13. Cấu hình SQL Server container là tùy chọn", 1)
    add_para(
        doc,
        "File docker-compose.yml hiện vẫn có service sqlserver dưới profile with-sqlserver để phục vụ trường hợp nhóm có backup riêng và muốn restore vào SQL Server container. Tuy nhiên, đây không nên là cách trình bày mặc định nếu người đọc chỉ mới làm quen với dự án."
    )
    add_code(doc, "docker compose --profile with-sqlserver up -d sqlserver")
    add_bullet(doc, "Chỉ nên dùng profile này khi nhóm đã có file backup .bak hoặc bộ script seed hoàn chỉnh.")
    add_bullet(doc, "Nếu không có dữ liệu nền đầy đủ, không nên dùng cấu hình SQL container cho buổi demo chính thức.")

    add_heading(doc, "14. Sự cố thường gặp và cách xử lý", 1)
    add_table(
        doc,
        ["Hiện tượng", "Nguyên nhân thường gặp", "Cách xử lý"],
        [
            ["docker compose up lỗi kết nối DB", "Sai mật khẩu SQL hoặc DB chưa tồn tại", "Kiểm tra lại .env, đăng nhập SQL bằng sqlcmd hoặc SSMS."],
            ["Web chạy nhưng không có dữ liệu", "Database Electronics_Shop chưa có schema hoặc chưa restore dữ liệu", "Restore DB mẫu hoặc dùng DB local đã kiểm thử."],
            ["recommend-api không lên healthy", "Không kết nối được SQL hoặc Redis", "Kiểm tra log container và trạng thái redis."],
            ["Không mở được localhost:8080", "Container web chưa chạy hoặc cổng bị chiếm", "Xem docker compose ps, đổi cổng hoặc giải phóng cổng 8080."],
            ["Migrate EF bị lỗi", "Schema đích chưa phù hợp với các migration hiện có", "Không migrate vào DB trắng; ưu tiên restore DB mẫu."],
        ],
    )

    add_heading(doc, "15. Ghi chú về định dạng tài liệu", 1)
    add_para(
        doc,
        "Tài liệu này được viết lại theo hướng dễ tiếp cận hơn cho sinh viên mới bắt đầu, đồng thời đã được sinh bằng Unicode tiếng Việt đầy đủ để tránh lỗi hiển thị ký tự."
    )
    add_para(
        doc,
        "Theo quy trình chuẩn, tài liệu DOCX nên được render kiểm tra trực quan trước khi bàn giao cuối cùng. Trong môi trường hiện tại, bước render vẫn phụ thuộc việc cài LibreOffice hoặc soffice."
    )

    add_heading(doc, "16. Kết luận", 1)
    add_para(
        doc,
        "Với cấu trúc hiện tại, Tech_Store đã được tổ chức theo hướng một bộ mã nguồn duy nhất để chạy đồng thời website ASP.NET Core MVC và Recommendation API viết bằng Python. Trong phạm vi đồ án sinh viên, cách trình bày phù hợp và trung thực nhất là mô tả mô hình đã kiểm thử được: Docker dùng cho các dịch vụ ứng dụng, còn SQL Server local trên máy host dùng để chứa database Electronics_Shop."
    )

    section = doc.add_section(WD_SECTION.NEW_PAGE)
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer.add_run("Tech_Store - Tài liệu triển khai và cài đặt")
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    run.font.size = Pt(10)

    OUTPUT_PATH.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT_PATH)
    print(OUTPUT_PATH)


if __name__ == "__main__":
    build_doc()
