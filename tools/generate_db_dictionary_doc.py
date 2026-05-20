import re
from collections import OrderedDict
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt


ROOT = Path(r"D:\Projects\Tech_Store")
SNAPSHOT = ROOT / "Migrations" / "ApplicationDbContextModelSnapshot.cs"
MODELS_DIR = ROOT / "Models"
OUTPUT_DIR = Path(r"D:\TaiLieuHocTap\MonHoc\Ky4\docx_render_review")
OUTPUT_FILE = OUTPUT_DIR / "Mo_ta_chi_tiet_35_bang_Tech_Store.docx"


TABLE_RE = re.compile(r'b\.ToTable\("([^"]+)"')
PROP_RE = re.compile(r'^\s*b\.Property<([^>]+)>\("([^"]+)"\)', re.MULTILINE)
COLNAME_RE = re.compile(r'\.HasColumnName\("([^"]+)"\)')
COLTYPE_RE = re.compile(r'\.HasColumnType\("([^"]+)"\)')
PK_RE = re.compile(r'b\.HasKey\((.*?)\)')
FK_RE = re.compile(
    r'b\.HasOne\("Tech_Store\.Models\.([^"]+)",\s*"[^"]*"\)\s*'
    r'(?:.|\n)*?\.HasForeignKey\("([^"]+)"\)',
    re.MULTILINE,
)


def normalize_type(raw_type: str, col_type: str | None) -> str:
    if col_type:
        return col_type
    mapping = {
        "int": "int",
        "int?": "int",
        "string": "nvarchar(max)",
        "bool": "bit",
        "bool?": "bit",
        "DateTime": "datetime",
        "DateTime?": "datetime",
        "decimal": "decimal(18,2)",
        "decimal?": "decimal(18,2)",
        "double": "float",
        "double?": "float",
        "float": "real",
        "float?": "real",
        "long": "bigint",
        "long?": "bigint",
    }
    return mapping.get(raw_type, raw_type)


def split_entities(snapshot_text: str):
    marker = 'modelBuilder.Entity("'
    parts = snapshot_text.split(marker)
    entities = []
    for chunk in parts[1:]:
        model_name = chunk.split('"', 1)[0]
        body = chunk.split("});", 1)[0]
        entities.append((model_name, body))
    return entities


def load_declared_properties():
    declared = {}
    prop_re = re.compile(r"public\s+[\w<>\?\., ]+\s+(\w+)\s*\{\s*get;\s*set;\s*\}")
    class_re = re.compile(r"public\s+partial\s+class\s+(\w+)")

    for path in MODELS_DIR.glob("*.cs"):
        text = path.read_text(encoding="utf-8")
        class_match = class_re.search(text)
        if not class_match:
            continue
        class_name = class_match.group(1)
        declared[class_name] = set(prop_re.findall(text))

    return declared


def reorder_columns(columns, pk_props):
    original_order = {col["prop_name"]: idx for idx, col in enumerate(columns)}
    address_rank = {
        "address_line": 0,
        "ward": 1,
        "district": 2,
        "province": 3,
    }
    trailing_rank = {
        "start_at": 0,
        "end_at": 1,
        "startedat": 0,
        "expiredat": 1,
        "payment_date": 0,
        "order_date": 0,
        "review_date": 0,
        "added_date": 0,
        "readat": 0,
        "created_verify": 0,
        "last_login": 0,
        "last_login_ip": 1,
        "last_login_device": 2,
        "last_request_at": 3,
        "last_request_ip": 4,
        "last_request_device": 5,
        "createdat": 98,
        "created_at": 98,
        "updatedat": 99,
        "updated_at": 99,
    }

    def sort_key(column):
        prop = column["prop_name"]
        lower = column["column_name"].lower()

        if prop in pk_props:
            return (0, original_order[prop])
        if lower.endswith("_id") or lower.endswith("id") or lower == "id":
            return (1, original_order[prop])
        if lower in {"code", "sku", "slug", "email", "phone_number"}:
            return (2, original_order[prop])
        if "name" in lower or lower in {"title"}:
            return (3, original_order[prop])
        if lower in address_rank:
            return (4, address_rank[lower])
        if lower in trailing_rank:
            return (9, trailing_rank[lower])
        return (5, original_order[prop])

    return sorted(columns, key=sort_key)


def parse_schema():
    text = SNAPSHOT.read_text(encoding="utf-8")
    entities = split_entities(text)
    declared_props = load_declared_properties()
    tables = OrderedDict()
    model_to_table = {}

    for model_name, body in entities:
        table_match = TABLE_RE.search(body)
        if not table_match:
            continue
        table_name = table_match.group(1)
        model_to_table[model_name.split(".")[-1]] = table_name
        if table_name in tables:
            continue

        model_short_name = model_name.split(".")[-1]
        allowed_props = declared_props.get(model_short_name, set())
        props = []
        prop_matches = list(PROP_RE.finditer(body))
        for idx, match in enumerate(prop_matches):
            raw_type, prop_name = match.groups()
            start = match.start()
            end = prop_matches[idx + 1].start() if idx + 1 < len(prop_matches) else len(body)
            prop_block = body[start:end]

            col_name_match = COLNAME_RE.search(prop_block)
            col_type_match = COLTYPE_RE.search(prop_block)
            if not col_name_match and prop_name not in allowed_props and table_name != "UserRole":
                continue
            col_name = col_name_match.group(1) if col_name_match else prop_name
            col_type = normalize_type(raw_type, col_type_match.group(1) if col_type_match else None)
            props.append(
                {
                    "prop_name": prop_name,
                    "column_name": col_name,
                    "data_type": col_type,
                }
            )

        pk_match = PK_RE.search(body)
        pk_props = set()
        if pk_match:
            pk_raw = pk_match.group(1)
            pk_props.update(re.findall(r'"([^"]+)"', pk_raw))

        tables[table_name] = {
            "model_name": model_short_name,
            "columns": reorder_columns(props, pk_props),
            "pk_props": pk_props,
            "fk_props": {},
        }

    for model_name, body in entities:
        short_model = model_name.split(".")[-1]
        table_name = model_to_table.get(short_model)
        if not table_name or table_name not in tables:
            continue
        for target_model, fk_prop in FK_RE.findall(body):
            target_table = model_to_table.get(target_model, target_model)
            tables[table_name]["fk_props"][fk_prop] = target_table

    return tables


def meaning_for_column(table_name: str, column: dict, pk_props: set[str], fk_props: dict[str, str]) -> str:
    prop_name = column["prop_name"]
    col_name = column["column_name"]
    lower = col_name.lower()

    if prop_name in pk_props:
        if len(pk_props) > 1:
            return f"Thành phần khóa chính ghép của bảng {table_name}"
        return f"Khóa chính của bảng {table_name}"

    if prop_name in fk_props:
        return f"Khóa ngoại liên kết đến bảng {fk_props[prop_name]}"

    if lower in {"created_at", "createdat"}:
        return "Thời điểm tạo bản ghi"
    if lower in {"updated_at", "updatedat"}:
        return "Thời điểm cập nhật gần nhất của bản ghi"
    if lower in {"name", "role_name", "title"}:
        return "Tên hiển thị của đối tượng"
    if lower == "description":
        return "Mô tả chi tiết của đối tượng"
    if "code" in lower and lower != "barcode":
        return "Mã định danh hoặc mã nghiệp vụ"
    if "email" in lower:
        return "Địa chỉ email"
    if "phone" in lower:
        return "Số điện thoại liên hệ"
    if "address" in lower:
        return "Thông tin địa chỉ"
    if "image" in lower or lower in {"img", "path"}:
        return "Đường dẫn hoặc hình ảnh minh họa"
    if "price" in lower or "amount" in lower:
        return "Giá trị tiền tệ"
    if lower == "quantity":
        return "Số lượng"
    if lower == "stock":
        return "Số lượng tồn kho"
    if "status" in lower:
        return "Trạng thái xử lý hoặc hoạt động"
    if lower.startswith("is_") or lower.startswith("is") or lower.startswith("visible") or lower.startswith("open_in"):
        return "Cờ điều khiển trạng thái hoặc phạm vi hiển thị"
    if "date" in lower or lower.endswith("_at") or lower.endswith("at"):
        return "Mốc thời gian liên quan đến nghiệp vụ"
    if lower in {"content", "message", "note", "notes", "comment", "metadata_json"}:
        return "Nội dung thông tin hoặc ghi chú"
    if lower in {"sku", "slug"}:
        return "Mã hoặc chuỗi định danh phục vụ tra cứu"
    if lower in {"type", "target_type", "payment_method", "event_type", "promotion"}:
        return "Loại hoặc phân loại dữ liệu"
    if lower in {"value", "weight", "rating", "priority", "sort_order"}:
        return "Giá trị dùng để tính toán, sắp xếp hoặc đánh giá"
    if lower == "parent_id":
        return f"Khóa tự tham chiếu thể hiện quan hệ cha - con trong bảng {table_name}"
    if lower == "categoryid1":
        return "Cột tự tham chiếu do EF sinh ra cho quan hệ phân cấp của Category"
    if lower in {"first_name", "last_name"}:
        return "Thành phần họ tên người dùng"
    if lower in {"province", "district", "ward"}:
        return "Đơn vị hành chính của địa chỉ"
    if lower == "session_id":
        return "Mã phiên làm việc của người dùng hoặc khách"
    if lower == "external_url" or lower == "redirecturl":
        return "Đường dẫn điều hướng"
    if lower == "product_sys_id":
        return "Mã hệ thống nội bộ của sản phẩm"

    return f"Thuộc tính {col_name} của bảng {table_name}"


def set_cell_text(cell, text, bold=False):
    cell.text = ""
    p = cell.paragraphs[0]
    run = p.add_run(text)
    run.bold = bold
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    run.font.size = Pt(13)


def format_table_borders(table):
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for border_name in ("top", "left", "bottom", "right", "insideH", "insideV"):
        border = borders.find(qn(f"w:{border_name}"))
        if border is None:
            border = OxmlElement(f"w:{border_name}")
            borders.append(border)
        border.set(qn("w:val"), "single")
        border.set(qn("w:sz"), "4")
        border.set(qn("w:space"), "0")
        border.set(qn("w:color"), "000000")


def apply_global_font(doc: Document):
    styles = doc.styles
    for style_name in ["Normal", "Title", "Heading 1", "Heading 2", "Heading 3"]:
        style = styles[style_name]
        style.font.name = "Times New Roman"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
        style.font.size = Pt(13)


def build_doc(tables):
    doc = Document()
    apply_global_font(doc)

    section = doc.sections[0]
    section.top_margin = Pt(72)
    section.bottom_margin = Pt(72)
    section.left_margin = Pt(72)
    section.right_margin = Pt(72)

    title = doc.add_paragraph()
    title.style = doc.styles["Title"]
    title_run = title.add_run("MÔ TẢ CẤU TRÚC CƠ SỞ DỮ LIỆU TECH_STORE")
    title_run.bold = True
    title_run.font.name = "Times New Roman"
    title_run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    title_run.font.size = Pt(13)

    intro = doc.add_paragraph(
        "Tài liệu mô tả đầy đủ 35 bảng dữ liệu hiện có trong hệ thống, bao gồm toàn bộ cột, kiểu dữ liệu và ý nghĩa nghiệp vụ của từng thuộc tính."
    )
    for run in intro.runs:
        run.font.name = "Times New Roman"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
        run.font.size = Pt(13)

    for idx, (table_name, info) in enumerate(tables.items(), start=1):
        if idx > 1:
            doc.add_paragraph()

        heading = doc.add_paragraph()
        heading.style = doc.styles["Heading 1"]
        heading_run = heading.add_run(f"{idx}. Bảng {table_name}")
        heading_run.bold = True
        heading_run.font.name = "Times New Roman"
        heading_run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
        heading_run.font.size = Pt(13)

        table = doc.add_table(rows=1, cols=4)
        table.style = "Table Grid"
        table.autofit = True

        headers = ["STT", "Tên thuộc tính", "Kiểu dữ liệu", "Ý nghĩa"]
        for i, header in enumerate(headers):
            set_cell_text(table.rows[0].cells[i], header, bold=True)

        for col_idx, column in enumerate(info["columns"], start=1):
            row = table.add_row().cells
            set_cell_text(row[0], str(col_idx))
            set_cell_text(row[1], column["column_name"])
            set_cell_text(row[2], column["data_type"])
            set_cell_text(
                row[3],
                meaning_for_column(
                    table_name,
                    column,
                    info["pk_props"],
                    info["fk_props"],
                ),
            )

        format_table_borders(table)

    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT_FILE)


def main():
    tables = parse_schema()
    expected_count = 35
    if len(tables) != expected_count:
        raise RuntimeError(f"Expected {expected_count} tables, found {len(tables)}")
    build_doc(tables)
    print(OUTPUT_FILE)


if __name__ == "__main__":
    main()
